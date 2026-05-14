"""
cv_parser.py — Extrae campos del perfil desde un CV en PDF o DOCX.
Devuelve un dict con los campos detectados (sin garantía de completitud).
"""
import os
import re


# ---------------------------------------------------------------------------
# Extracción de texto
# ---------------------------------------------------------------------------

def _text_from_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            pages.append(t)
    return "\n".join(pages)


def _text_from_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # También tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _text_from_pdf(path)
    elif ext in (".docx", ".doc"):
        return _text_from_docx(path)
    raise ValueError(f"Formato no soportado: {ext}. Usa PDF o DOCX.")


# ---------------------------------------------------------------------------
# Extracción de campos
# ---------------------------------------------------------------------------

def _find_email(text: str) -> str | None:
    m = re.search(r"[\w.+\-]+@[\w\-]+\.[a-z]{2,}", text, re.IGNORECASE)
    return m.group(0).lower() if m else None


def _find_phone(text: str) -> tuple[str | None, str | None]:
    """Retorna (phone_formatted, phone_digits_only)"""
    # Formato chileno: +56 9 XXXX XXXX  o  +569XXXXXXXX  o  9XXXXXXXX
    m = re.search(r"(\+56[\s\-]?)?9[\s\-]?\d{4}[\s\-]?\d{4}", text)
    if not m:
        return None, None
    raw = m.group(0).strip()
    digits = re.sub(r"[^\d]", "", raw)
    # Normalizar: asegurar que empiece con 56
    if not digits.startswith("56"):
        digits = "56" + digits
    local = digits[2:]          # 9XXXXXXXX
    formatted = f"+56 {local[0]} {local[1:5]} {local[5:]}"
    return formatted, local


def _find_linkedin(text: str) -> str | None:
    m = re.search(r"linkedin\.com/in/[\w\-]+", text, re.IGNORECASE)
    if not m:
        return None
    url = m.group(0)
    return f"https://www.{url}" if not url.startswith("http") else url


def _find_portfolio(text: str) -> str | None:
    m = re.search(r"github\.com/[\w\-]+(?:/[\w\-]+)?", text, re.IGNORECASE)
    if not m:
        return None
    return f"https://{m.group(0)}"


def _find_name(lines: list[str]) -> tuple[str | None, str | None, str | None]:
    """full_name, first_name, last_name"""
    for line in lines[:8]:
        words = line.split()
        if 2 <= len(words) <= 5:
            # Todos en mayúsculas (nombre en cabecera) O title case
            if all(w.replace("Á","A").replace("É","E").replace("Í","I")
                     .replace("Ó","O").replace("Ú","U").isupper() for w in words):
                full = " ".join(w.title() for w in words)
                return full, words[0].title(), " ".join(words[1:]).title()
            if all(w[0].isupper() for w in words if w and w[0].isalpha()):
                # Evitar líneas que son títulos de sección (PERFIL, EDUCACIÓN…)
                if not any(kw in line.upper() for kw in
                           ["PERFIL", "EXPERIENCIA", "EDUCACIÓN", "EDUCACION",
                            "HABILIDADES", "IDIOMAS", "COMPETENCIAS"]):
                    return line, words[0], " ".join(words[1:])
    return None, None, None


def _find_city(text: str) -> str | None:
    cities = [
        "Maipú","Maipu","Santiago","Providencia","Las Condes","Ñuñoa","Nuñoa",
        "Vitacura","La Florida","Pudahuel","Quilicura","Peñalolén","San Bernardo",
        "Puente Alto","El Bosque","La Pintana","Cerrillos","Renca","Conchalí",
        "Estación Central","Lo Barnechea","Huechuraba","Recoleta","Independencia",
        "Concepción","Valparaíso","Viña del Mar","Antofagasta","Temuco",
    ]
    for city in cities:
        if re.search(rf"\b{re.escape(city)}\b", text, re.IGNORECASE):
            # Intenta capturar "Ciudad, RM" o "Ciudad, Región …"
            m = re.search(
                rf"({re.escape(city)}[^,\n]{{0,40}}(?:RM|Región Metropolitana)?)",
                text, re.IGNORECASE
            )
            if m:
                return m.group(1).strip().rstrip(",").strip()
            return city
    return None


def _find_cover_letter(text: str) -> str | None:
    """Párrafo bajo PERFIL PROFESIONAL o SOBRE MÍ."""
    m = re.search(
        r"(?:PERFIL\s+PROFESIONAL|SOBRE\s+M[IÍ]|RESUMEN\s+PROFESIONAL|SUMMARY)"
        r"[:\s\n]+(.+?)(?:\n\n|\n[A-ZÁÉÍÓÚ]{3,}\s)",
        text, re.IGNORECASE | re.DOTALL
    )
    if m:
        txt = " ".join(m.group(1).split())
        return txt[:600]
    return None


def _find_years_exp(text: str) -> str:
    """Calcula años de experiencia laboral a partir de fechas encontradas."""
    # Busca patrones como "Nov 2020" o "2020" en contexto de experiencia
    years = re.findall(r"\b(20\d{2})\b", text)
    if len(years) >= 2:
        years_int = [int(y) for y in years]
        span = max(years_int) - min(years_int)
        return str(span) if span > 0 else "0"
    return "0"


# ---------------------------------------------------------------------------
# Función principal
# ---------------------------------------------------------------------------

def parse_cv(path: str) -> dict:
    """
    Parsea un CV (PDF o DOCX) y retorna un dict con los campos encontrados.
    Solo incluye claves cuyo valor fue detectado (no None).
    Claves retornadas (si se encuentran):
      full_name, first_name, last_name, email, phone, phone_number,
      city, linkedin, portfolio, cover_letter, years_exp
    """
    text = extract_text(path)
    lines = [l.strip() for l in text.splitlines() if l.strip()]

    result: dict[str, str] = {}

    email = _find_email(text)
    if email:
        result["email"] = email

    phone, phone_num = _find_phone(text)
    if phone:
        result["phone"] = phone
    if phone_num:
        result["phone_number"] = phone_num

    linkedin = _find_linkedin(text)
    if linkedin:
        result["linkedin"] = linkedin

    portfolio = _find_portfolio(text)
    if portfolio:
        result["portfolio"] = portfolio

    full_name, first_name, last_name = _find_name(lines)
    if full_name:
        result["full_name"] = full_name
    if first_name:
        result["first_name"] = first_name
    if last_name:
        result["last_name"] = last_name

    city = _find_city(text)
    if city:
        result["city"] = city

    cover = _find_cover_letter(text)
    if cover:
        result["cover_letter"] = cover

    result["years_exp"] = _find_years_exp(text)

    return result
